import os
import re
import json
import docx
from datetime import datetime
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable

from db.session import SessionLocal, init_db
from db.models import SynopsisVersion

HUMANIZED_TEXT_PATH = r"E:\Research\Draft\Latest Drafts\synopsis_humanized_text.txt"
SYNOPSIS_MD_PATH = r"E:\Research\Synopsis_With_Radiologist_Comparison.md"

def get_synopsis_text():
    """Retrieve raw synopsis text from database (active version), disk file, or fallback."""
    init_db()
    db = SessionLocal()
    try:
        active_v = db.query(SynopsisVersion).filter(SynopsisVersion.is_active == 1).order_by(SynopsisVersion.id.desc()).first()
        if active_v and active_v.content:
            return active_v.content
    except Exception as e:
        print(f"DB fetch warning: {e}")
    finally:
        db.close()

    if os.path.exists(HUMANIZED_TEXT_PATH):
        with open(HUMANIZED_TEXT_PATH, "r", encoding="utf-8") as f:
            return f.read()
    elif os.path.exists(SYNOPSIS_MD_PATH):
        with open(SYNOPSIS_MD_PATH, "r", encoding="utf-8") as f:
            return f.read()
    else:
        return "# Research Synopsis\n\nNo synopsis file found on disk."

def save_synopsis_text(text: str, version_tag: str = "Draft Update", author: str = "Dr. Muhammad Mudassir"):
    """Save updated synopsis text to database AND disk files for total persistence."""
    init_db()
    db = SessionLocal()
    try:
        # Mark previous active version inactive
        db.query(SynopsisVersion).update({"is_active": 0})
        
        word_count = len(text.split())
        new_v = SynopsisVersion(
            version_tag=version_tag,
            content=text,
            word_count=word_count,
            author=author,
            is_active=1
        )
        db.add(new_v)
        db.commit()
    except Exception as e:
        print(f"DB save warning: {e}")
        db.rollback()
    finally:
        db.close()

    # Save to disk files
    os.makedirs(os.path.dirname(HUMANIZED_TEXT_PATH), exist_ok=True)
    with open(HUMANIZED_TEXT_PATH, "w", encoding="utf-8") as f:
        f.write(text)
    
    with open(SYNOPSIS_MD_PATH, "w", encoding="utf-8") as f:
        f.write(text)
    return True

def list_synopsis_versions():
    """List all saved synopsis versions from database."""
    init_db()
    db = SessionLocal()
    try:
        versions = db.query(SynopsisVersion).order_by(SynopsisVersion.id.desc()).all()
        return [
            {
                "id": v.id,
                "version_tag": v.version_tag,
                "word_count": v.word_count,
                "author": v.author,
                "is_active": bool(v.is_active),
                "created_at": v.created_at.strftime("%Y-%m-%d %H:%M:%S") if v.created_at else "N/A"
            }
            for v in versions
        ]
    finally:
        db.close()

def restore_synopsis_version(version_id: int):
    """Restore a specific version snapshot as active."""
    init_db()
    db = SessionLocal()
    try:
        target_v = db.query(SynopsisVersion).filter(SynopsisVersion.id == version_id).first()
        if not target_v:
            return False
        
        db.query(SynopsisVersion).update({"is_active": 0})
        target_v.is_active = 1
        db.commit()
        
        # Save to disk files
        with open(HUMANIZED_TEXT_PATH, "w", encoding="utf-8") as f:
            f.write(target_v.content)
        with open(SYNOPSIS_MD_PATH, "w", encoding="utf-8") as f:
            f.write(target_v.content)
        return True
    finally:
        db.close()

def generate_synopsis_docx(output_path: str):
    """Generate UHS-compliant Word (.docx) document from synopsis text."""
    text = get_synopsis_text()
    
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    
    lines = text.split("\n")
    for line in lines:
        line_str = line.strip()
        if not line_str or line_str.startswith("---"):
            continue
            
        if line_str.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line_str[2:])
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line_str.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line_str[3:])
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif line_str.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line_str[4:])
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif line_str.startswith("**") and line_str.endswith("**"):
            p = doc.add_paragraph()
            clean = line_str.replace("**", "")
            run = p.add_run(clean)
            run.bold = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
        else:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*)", line_str)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
            
    doc.save(output_path)
    return output_path

def generate_synopsis_pdf(output_path: str):
    """Generate printable PDF document using ReportLab."""
    text = get_synopsis_text()
    
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#1E3A8A'),
        spaceAfter=12,
        alignment=1
    )
    
    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#1E3A8A'),
        spaceBefore=14,
        spaceAfter=6
    )

    h3_style = ParagraphStyle(
        'DocH3',
        parent=styles['Heading3'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor('#374151'),
        spaceBefore=10,
        spaceAfter=4
    )
    
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=11,
        leading=15,
        textColor=colors.HexColor('#1F2937'),
        spaceAfter=8
    )

    story = []
    lines = text.split("\n")
    
    for line in lines:
        line_str = line.strip()
        if not line_str or line_str.startswith("---"):
            continue
            
        if line_str.startswith("# "):
            story.append(Paragraph(line_str[2:], title_style))
            story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#1E3A8A'), spaceBefore=4, spaceAfter=12))
        elif line_str.startswith("## "):
            story.append(Paragraph(line_str[3:], h2_style))
        elif line_str.startswith("### "):
            story.append(Paragraph(line_str[4:], h3_style))
        else:
            formatted = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", line_str)
            story.append(Paragraph(formatted, body_style))
            
    doc.build(story)
    return output_path
